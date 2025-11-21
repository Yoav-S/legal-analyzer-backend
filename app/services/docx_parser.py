"""
DOCX text extraction service.
"""
import io
from docx import Document

from app.utils.logger import setup_logger
from app.utils.errors import ProcessingError

logger = setup_logger(__name__)


class DOCXParser:
    """Service for extracting text from DOCX files."""
    
    @staticmethod
    async def extract_text(docx_content: bytes) -> str:
        """
        Extract text from DOCX content.
        
        Args:
            docx_content: DOCX file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            docx_file = io.BytesIO(docx_content)
            doc = Document(docx_file)
            
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise ProcessingError(f"Failed to extract text from DOCX: {str(e)}")

